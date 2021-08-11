import time

import docx
import os
import unicodedata
from docx2pdf import convert
from openpyxl import Workbook, load_workbook


def importExcel(fileroute):
    #TODO la ruta de la base no puede ser estatia
    if not fileroute:  # base
        fileroute = "C:/Users/Javier/Documents/MyCodeProjects/CAG/FilexMaker/documents/alumnosSampleFilex.xlsx"

    try:
        base = load_workbook(fileroute)
        baseHoja = base.active
        print(fileroute)
    except:
        print("ERROR IMPORTING EXCEL")
        return

    return baseHoja

def importMachote(docsroute):
    # RECORRIENDO EL ARCHIVO EN BUSCA DE SUSTITUCION
    if not docsroute:
        machote = docx.Document("DOCUMENTS/machotefilex.docx")
    try:
        #machote = docx.Document(docsroute)
        pass
        #print(doc)
    except:
        print("ERROR IMPORTANDO EL MACHOTE DOCX")

    return machote

def getNombreBase(hojaExcel, pos):
    try:
        nombre = unicodedata.normalize('NFD', ((hojaExcel['C' + str(pos)].value)))
    except:
        nombre = "-"
    try:
        apellidoPaterno = unicodedata.normalize('NFD', ((hojaExcel['D' + str(pos)].value)))
    except:
        apellidoPaterno = "-"
    try:
        apellidoMaterno = unicodedata.normalize('NFD', ((hojaExcel['E' + str(pos)].value)))
    except:
        apellidoMaterno = "-"

    fullName = [nombre, apellidoPaterno, apellidoMaterno]
    return fullName;

def getNombreFilex(hojaExcel,pos):
    try:
        fullNombre = unicodedata.normalize('NFD', ((hojaExcel['G' + str(pos)].value)))
    except:
        fullNombre =""
    print(fullNombre)
    return fullNombre

def getNivelacion(hojaExcel, pos):
    try:
        nivel =  str(hojaExcel['D' + str(pos)].value)
    except:
        nivel = ""
    return nivel

def getHoras(hojaExcel, pos):
    try:
        horas = unicodedata.normalize('NFD', ((hojaExcel['B' + str(pos)].value)))
    except:
        horas = ""
    return horas

def getDias(hojaExcel,pos):
    try:
        dias = unicodedata.normalize('NFD', ((hojaExcel['A' + str(pos)].value)))
        print("DIAS|",dias)
    except:
        dias = ""
    return dias

def getCiclo(hojaExcel,pos):
    try:
        ciclo = unicodedata.normalize('NFD', ((hojaExcel['F' + str(pos)].value)))
    except:
        ciclo = ""
    return ciclo

def isEmpty(hojaExcel,pos):
        #nivel
    if (hojaExcel['D' + str(pos)].value) is None and (hojaExcel['G' + str(pos)].value) is None:
        return True
    else:
        return False


def scanDocument(i,machoteRoute,excelRoute,destino,num):
    # RECORRIENDO EL ARCHIVO EN BUSCA DE SUSTITUCION
    print(i)

    print(excelRoute)

    if not destino:
        destino ="C:/Users/Javier/Documents/MyCodeProjects/CAG/FilexMaker/documents/generadas"


    if(isEmpty(excelRoute,i)):
        return True
        print("termino")
    else:
        pass

    for paragraph in machoteRoute.paragraphs:
        #print(paragraph.text)
        font = paragraph.style.font
        font.name = 'Arial'
        if ("$DIAS$") in paragraph.text:
            dias = getDias(excelRoute, i)
            paragraph.text = paragraph.text.replace("$DIAS$", dias)
        if ("$GENERO$") in paragraph.text:
            # PENDIENTE
            gender = "el"
            paragraph.text = paragraph.text.replace("$GENERO$", gender)
        if ("$NOMBRE$") in paragraph.text:
            namef = getNombreFilex(excelRoute, i)
            paragraph.text = paragraph.text.replace("$NOMBRE$", namef)
        if ("$NIVELACION$") in paragraph.text:
            nivelf = getNivelacion(excelRoute, i)
            paragraph.text = paragraph.text.replace("$NIVELACION$", nivelf)
        if ("$HORAS$") in paragraph.text:
            horas = getHoras(excelRoute, i)
            paragraph.text = paragraph.text.replace("$HORAS$", horas)
        if ("$CICLO$") in paragraph.text:
            ciclo = getCiclo(excelRoute, i)
            paragraph.text = paragraph.text.replace("$CICLO$", ciclo)
        if ("$NUMERO$") in paragraph.text:
            numero = str(num)
            paragraph.text = paragraph.text.replace("$NUMERO$", numero)

        # print(paragraph.text)
        # print("---")


    #destino ="C:/Users/Javier/Documents/MyCodeProjects/CAG/FilexMaker/documents/generadas"
    machoteRoute.save(destino+"/temp" + str(i) + ".docx")
    #time.sleep(0.3)

    convert(destino+"/temp" + str(i) + ".docx", destino+"/"+ namef +"_NIVEL_"+nivelf+ ".pdf")
    #time.sleep(0.3)
    os.remove(destino+"/temp"+str(i)+".docx")
    #time.sleep(0.3)
    print("finalizado")

    time.sleep(1)
    return False

    #





# Press the green button in the gutter to run the script.
if __name__ == '__main__':


    for i in range(2,100):
        if scanDocument(i, importMachote(""), importExcel(""), "", "1") == True:
            break

        print("iterador",i)

    #scanDocument(2, importMachote(""), importExcel(""), "", "1")
    #scanDocument(3,importMachote(""),importExcel(""),"","1")
    #scanDocument(i,machoteRoute,excelRoute,destino,num):



